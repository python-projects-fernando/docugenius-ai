import json
import logging
from typing import Dict, Any
from docx import Document
from io import BytesIO
from backend.application.repositories.document_type_repository import DocumentTypeRepository
from backend.application.repositories.document_field_repository import DocumentFieldRepository
from backend.application.repositories.generated_document_repository import GeneratedDocumentRepository
from backend.application.ai_gateway.ai_gateway import AIGateway
from backend.application.file_storage.file_storage import FileStorageGateway
from backend.core.models.document_type import DocumentType as CoreDocumentType
from backend.core.models.document_field import DocumentField as CoreDocumentField
from backend.core.models.generated_document import GeneratedDocument as CoreGeneratedDocument
from backend.application.dtos.document_generation import GenerateDocumentRequest
from backend.application.dtos.api_response import APIResponse
from backend.application.prompts import GENERATE_DOCUMENT_CONTENT_PROMPT

logger = logging.getLogger(__name__)

class GenerateDocumentUseCase:
    def __init__(
        self,
        document_type_repo: DocumentTypeRepository,
        document_field_repo: DocumentFieldRepository,
        generated_document_repo: GeneratedDocumentRepository,
        ai_gateway: AIGateway,
        file_storage_gateway: FileStorageGateway
    ):
        self._document_type_repo = document_type_repo
        self._document_field_repo = document_field_repo
        self._generated_document_repo = generated_document_repo
        self._ai_gateway = ai_gateway
        self._file_storage_gateway = file_storage_gateway

    async def execute(self, request_dto: GenerateDocumentRequest, current_user_id: int) -> APIResponse[dict]:
        try:
            document_type_entity: CoreDocumentType = await self._document_type_repo.find_by_id(request_dto.document_type_id)
            if not document_type_entity:
                return APIResponse[dict](
                    success=False,
                    message=f"DocumentType with ID {request_dto.document_type_id} not found.",
                    error_code="DOC_TYPE_NOT_FOUND",
                    errors=[f"Cannot generate document: DocumentType with ID {request_dto.document_type_id} does not exist."],
                    data=None
                )

            fields_for_doc_type: list[CoreDocumentField] = await self._document_field_repo.find_all_by_document_type(request_dto.document_type_id)

            required_fields_missing = []
            for field_def in fields_for_doc_type:
                if field_def.is_required and field_def.name not in request_dto.filled_fields:
                    required_fields_missing.append(field_def.name)

            if required_fields_missing:
                return APIResponse[dict](
                    success=False,
                    message="Required fields are missing for document generation.",
                    error_code="MISSING_REQUIRED_FIELDS",
                    errors=[f"Field '{field_name}' is required but was not provided." for field_name in required_fields_missing],
                    data=None
                )

            filled_fields_json_str = json.dumps(request_dto.filled_fields, indent=2, ensure_ascii=False)
            prompt = GENERATE_DOCUMENT_CONTENT_PROMPT.format(
                document_type_name=document_type_entity.name,
                document_type_description=document_type_entity.description,
                filled_fields_json=filled_fields_json_str
            )

            from backend.application.dtos.ai_inference import InferenceRequest
            ai_request = InferenceRequest(
                model="meta-llama/Llama-3.1-8B-Instruct:cerebras",
                prompt=prompt
            )

            ai_response = await self._ai_gateway.generate_text(ai_request)
            generated_content = ai_response.generated_text

            import uuid
            unique_filename = f"generated_doc_{request_dto.document_type_id}_{uuid.uuid4().hex}.docx"

            doc = Document()
            doc.add_paragraph(generated_content)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            location_identifier = await self._file_storage_gateway.save_document(
                content=buffer.getvalue(),
                filename=unique_filename
            )

            generated_doc_entity = CoreGeneratedDocument(
                id=None,
                user_id=current_user_id,
                document_type_id=request_dto.document_type_id,
                file_path_or_key=location_identifier,
            )

            saved_entity = await self._generated_document_repo.save(generated_doc_entity)
            download_url = await self._file_storage_gateway.get_file_url(location_identifier)

            return APIResponse[dict](
                success=True,
                message="Document generated successfully by AI, saved using the configured storage gateway, and record stored in database.",
                data={
                    "location_identifier": saved_entity.file_path_or_key,
                    "download_url": download_url
                },
                error_code=None,
                errors=None
            )

        except Exception as e:
            logger.error(f"Error during document generation: {e}")
            return APIResponse[dict](
                success=False,
                message="An unexpected error occurred during document generation.",
                error_code="GENERATE_DOC_ERROR",
                errors=[f"Internal error: {str(e)}"],
                data=None
            )


# import json
# import logging
# import os
# from io import BytesIO
# from pathlib import Path
# from typing import Dict, Any
# from docx import Document
# from backend.application.repositories.document_type_repository import DocumentTypeRepository
# from backend.application.repositories.document_field_repository import DocumentFieldRepository
# from backend.application.ai_gateway.ai_gateway import AIGateway
# from backend.application.file_storage.file_storage import FileStorageGateway
# from backend.core.models.document_type import DocumentType as CoreDocumentType
# from backend.core.models.document_field import DocumentField as CoreDocumentField
# from backend.application.dtos.document_generation import GenerateDocumentRequest
# from backend.application.dtos.api_response import APIResponse
# from backend.application.prompts import GENERATE_DOCUMENT_CONTENT_PROMPT
#
# logger = logging.getLogger(__name__)
#
# class GenerateDocumentUseCase:
#     def __init__(self, document_type_repo: DocumentTypeRepository, document_field_repo: DocumentFieldRepository,
#                  ai_gateway: AIGateway,file_storage: FileStorageGateway):
#         self._document_type_repo = document_type_repo
#         self._document_field_repo = document_field_repo
#         self._ai_gateway = ai_gateway
#         self._temp_docs_dir = Path("backend") / "temp_generated_docs"
#         self._temp_docs_dir.mkdir(parents=True, exist_ok=True)
#         self._file_storage = file_storage
#
#     async def execute(self, request_dto: GenerateDocumentRequest) -> APIResponse[dict]:
#         try:
#             document_type_entity: CoreDocumentType = await self._document_type_repo.find_by_id(request_dto.document_type_id)
#             if not document_type_entity:
#                 return APIResponse[dict](
#                     success=False,
#                     message=f"DocumentType with ID {request_dto.document_type_id} not found.",
#                     error_code="DOC_TYPE_NOT_FOUND",
#                     errors=[f"Cannot generate document: DocumentType with ID {request_dto.document_type_id} does not exist."],
#                     data=None
#                 )
#
#             fields_for_doc_type: list[CoreDocumentField] = await self._document_field_repo.find_all_by_document_type(request_dto.document_type_id)
#
#             required_fields_missing = []
#             for field_def in fields_for_doc_type:
#                 if field_def.is_required and field_def.name not in request_dto.filled_fields:
#                     required_fields_missing.append(field_def.name)
#
#             if required_fields_missing:
#                 return APIResponse[dict](
#                     success=False,
#                     message="Required fields are missing for document generation.",
#                     error_code="MISSING_REQUIRED_FIELDS",
#                     errors=[f"Field '{field_name}' is required but was not provided." for field_name in required_fields_missing],
#                     data=None
#                 )
#
#             filled_fields_json_str = json.dumps(request_dto.filled_fields, indent=2, ensure_ascii=False)
#             prompt = GENERATE_DOCUMENT_CONTENT_PROMPT.format(
#                 document_type_name=document_type_entity.name,
#                 document_type_description=document_type_entity.description,
#                 filled_fields_json=filled_fields_json_str
#             )
#
#             from backend.application.dtos.ai_inference import InferenceRequest
#             ai_request = InferenceRequest(
#                 model="meta-llama/Llama-3.1-8B-Instruct:cerebras",
#                 prompt=prompt
#             )
#
#             ai_response = await self._ai_gateway.generate_text(ai_request)
#             generated_content = ai_response.generated_text
#
#             import uuid
#             unique_filename = f"generated_doc_{request_dto.document_type_id}_{uuid.uuid4().hex}.docx"
#             full_file_path = self._temp_docs_dir / unique_filename
#
#             doc = Document()
#             doc.add_paragraph(generated_content)
#
#             buffer = BytesIO()
#             doc.save(buffer)
#             buffer.seek(0)
#
#             # doc.save(full_file_path)
#
#             location_identifier = await self._file_storage.save_document(
#                 content=buffer.getvalue(),
#                 filename=unique_filename
#             )
#
#             return APIResponse[dict](
#                 success=True,
#                 message="Document generated successfully by AI and saved using the configured storage gateway.",
#                 data={"location_identifier": location_identifier},
#                 error_code=None,
#                 errors=None
#             )
#
#         except Exception as e:
#             logger.error(f"Error during document generation: {e}")
#             return APIResponse[dict](
#                 success=False,
#                 message="An unexpected error occurred during document generation.",
#                 error_code="GENERATE_DOC_ERROR",
#                 errors=[f"Internal error: {str(e)}"],
#                 data=None
#             )